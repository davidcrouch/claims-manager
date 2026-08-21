"""Build invoice WH-2026-08-20: billable ledger after WH-2026-04-30 through 2026-08-06.

Line items omit per-entry date (commit, hours, metrics, summary, description only).

Vendor from ABR / ASIC (extracted 2026-08-20):
- Entity: MORE0 PTY LTD
- ABN: 86 701 251 587 (active from 01 Jul 2026)
- ACN: 701 251 587
- Main business location: NSW 2533
- ABR GST: not currently registered (invoice still shows GST lines —
  confirm registration before issuing as a tax invoice)

Bill to unchanged: Branlamie Pty Ltd.
Rate: $125.00/hr ex GST + 10% GST (pending GST registration check).
"""
from __future__ import annotations

import re
from decimal import Decimal
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

A = "Arial"
WORK_HOURS = Path(__file__).resolve().parents[1] / "work_hours.md"
OUT_DIR = Path(__file__).resolve().parent
RATE_H = Decimal("125.00")

VENDOR_NAME = "More0 Pty Ltd"
VENDOR_ABN = "86 701 251 587"
VENDOR_ACN = "701 251 587"
VENDOR_BSB = "062-562"
VENDOR_ACCOUNT = "1035-8894"

# Inclusive upper bound; exclusive lower bound is prior invoice date.
AFTER_INVOICE_DATE = "2026-04-30"
PERIOD_END = "2026-08-06"

SPEC = {
    "date": "2026-08-20",
    "number": "WH-2026-08-20",
    "filename": "claims-manager-work-hours-invoice-2026-08-20.docx",
    "period": (
        "Period: new billable ledger entries after invoice WH-2026-04-30 "
        "through 2026-08-06 inclusive "
        "(ledger dates 2026-05-13 through 2026-08-02; no entries dated 2026-08-03–06)."
    ),
    "prior": "WH-2026-04-30",
}


def parse_ledger(path: Path) -> list[dict]:
    content = path.read_text(encoding="utf-8")
    header_re = re.compile(
        r"^- `(?P<date>\d{4}-\d{2}-\d{2})` `(?P<commit>[a-f0-9]+)` \*\*(?P<hours>[\d.]+) h\*\*(?P<suffix>[^\n]*)",
        re.MULTILINE,
    )
    matches = list(header_re.finditer(content))
    out: list[dict] = []
    for i, m in enumerate(matches):
        if "NOT-BILLABLE" in m.group("suffix"):
            continue
        hours_n = Decimal(m.group("hours"))
        if hours_n <= 0:
            continue
        date = m.group("date")
        if date <= AFTER_INVOICE_DATE or date > PERIOD_END:
            continue
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(content)
        block = content[start:end]
        lines = block.strip().split("\n")
        metrics = ""
        idx = 0
        if lines and lines[0].lstrip().startswith("`"):
            metrics = lines[0].strip().strip("`").strip()
            idx = 1
        lay = ""
        while idx < len(lines):
            s = lines[idx].strip()
            if s.startswith("Lay summary:"):
                lay = s.split("Lay summary:", 1)[1].strip()
                idx += 1
                break
            idx += 1
        desc: list[str] = []
        while idx < len(lines):
            raw = lines[idx].rstrip()
            if raw.strip().startswith("- `") and "` `" in raw and "**" in raw and " h**" in raw:
                break
            if raw.strip():
                desc.append(raw)
            idx += 1
        out.append(
            {
                "date": date,
                "commit": m.group("commit"),
                "hours": m.group("hours") + " h",
                "hours_n": hours_n,
                "metrics": metrics,
                "lay": lay,
                "desc": desc,
            }
        )
    return out


def add_runs_from_md(paragraph, text: str, size_pt: float) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        body = part[2:-2] if bold else part
        r = paragraph.add_run(body)
        r.bold = bold
        r.font.name = A
        r.font.size = Pt(size_pt)


def fmt_aud(d: Decimal) -> str:
    return f"${d:,.2f}"


def build_invoice(spec: dict, entries: list[dict]) -> Path:
    if not entries:
        raise SystemExit(
            f"{spec['number']}: no ledger entries after {AFTER_INVOICE_DATE} through {PERIOD_END}"
        )

    total_h = sum((e["hours_n"] for e in entries), Decimal("0"))
    sub_ex = total_h * RATE_H
    gst = (sub_ex * Decimal("0.10")).quantize(Decimal("0.01"))
    total_inc = sub_ex + gst

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    ht = doc.add_table(rows=2, cols=2)
    c00, c01 = ht.rows[0].cells[0], ht.rows[0].cells[1]
    c10, c11 = ht.rows[1].cells[0], ht.rows[1].cells[1]
    c10.merge(c11)

    p = c00.paragraphs[0]
    r = p.add_run(VENDOR_NAME)
    r.bold = True
    r.font.name = A
    r.font.size = Pt(14)
    p = c00.add_paragraph()
    r = p.add_run(f"ABN: {VENDOR_ABN}")
    r.font.name = A
    r.font.size = Pt(10)
    p = c00.add_paragraph()
    r = p.add_run(f"ACN: {VENDOR_ACN}")
    r.font.name = A
    r.font.size = Pt(10)

    p = c01.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("INVOICE")
    r.bold = True
    r.font.name = A
    r.font.size = Pt(22)
    p = c01.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"Invoice date: {spec['date']}")
    r.font.name = A
    r.font.size = Pt(10)
    p = c01.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"Invoice no.: {spec['number']}")
    r.font.name = A
    r.font.size = Pt(10)

    p = c10.paragraphs[0]
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run("Bill to")
    r.bold = True
    r.font.name = A
    r.font.size = Pt(10)
    for line in ("Branlamie Pty Ltd", "ABN 91 117 499 206", "ATTN: Brett Elsey"):
        p = c10.add_paragraph()
        r = p.add_run(line)
        r.font.name = A
        r.font.size = Pt(10)

    doc.add_paragraph()
    h = doc.add_heading("Work hours line items", level=1)
    for run in h.runs:
        run.font.name = A

    tbl = doc.add_table(rows=1 + len(entries) + 1, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for j, title in enumerate(("Entry", "Description (invoice text)")):
        rr = hdr[j].paragraphs[0].add_run(title)
        rr.bold = True
        rr.font.name = A
        rr.font.size = Pt(10)

    for i, e in enumerate(entries):
        row = tbl.rows[i + 1].cells
        left, right = row[0], row[1]
        left.text = ""
        lp = left.paragraphs[0]
        r = lp.add_run("-")
        r.font.name = A
        r.font.size = Pt(9)

        def left_line(label: str, value: str) -> None:
            p = left.add_paragraph()
            lr = p.add_run(f"{label}: ")
            lr.bold = True
            lr.font.name = A
            lr.font.size = Pt(9)
            vr = p.add_run(value)
            vr.font.name = A
            vr.font.size = Pt(9)

        # No per-entry date — commit, hours, metrics, summary only.
        left_line("Commit", e["commit"])
        left_line("Hours", e["hours"])
        left_line("Metrics", e["metrics"])
        sp = left.add_paragraph()
        sp.paragraph_format.space_before = Pt(6)
        sr = sp.add_run("Summary: ")
        sr.bold = True
        sr.font.name = A
        sr.font.size = Pt(9)
        sr2 = sp.add_run(e["lay"])
        sr2.font.name = A
        sr2.font.size = Pt(9)

        right.text = ""
        for j, line in enumerate(e["desc"]):
            rp = right.paragraphs[0] if j == 0 else right.add_paragraph()
            rp.paragraph_format.space_after = Pt(0)
            add_runs_from_md(rp, line, 9)

    tot = tbl.rows[-1].cells
    tot[0].text = ""
    tp = tot[0].paragraphs[0]
    tr = tp.add_run(f"Total (billable): {total_h} h")
    tr.bold = True
    tr.font.name = A
    tr.font.size = Pt(10)
    tot[1].text = ""
    tp2 = tot[1].paragraphs[0]
    tr2 = tp2.add_run(f"{len(entries)} line items")
    tr2.bold = True
    tr2.font.name = A
    tr2.font.size = Pt(10)

    doc.add_paragraph()
    h3 = doc.add_heading("Amount & payment", level=2)
    for run in h3.runs:
        run.font.name = A

    amount_lines = [
        (f"Billable hours: {total_h} h", False, 10),
        ("Rate: $125.00/hr (ex GST)", False, 10),
        (f"Subtotal (ex GST): {fmt_aud(sub_ex)}", False, 10),
        (f"GST (10%): {fmt_aud(gst)}", False, 10),
        (f"Total payable (inc GST): {fmt_aud(total_inc)}", True, 11),
    ]
    for text, bold, sz in amount_lines:
        p = doc.add_paragraph()
        rr = p.add_run(text)
        rr.bold = bold
        rr.font.name = A
        rr.font.size = Pt(sz)

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Pay to: ")
    r.bold = True
    r.font.name = A
    r.font.size = Pt(10)
    r2 = p.add_run(VENDOR_NAME)
    r2.font.name = A
    r2.font.size = Pt(10)
    for line in (f"BSB: {VENDOR_BSB}", f"Account: {VENDOR_ACCOUNT}"):
        p = doc.add_paragraph()
        rr = p.add_run(line)
        rr.font.name = A
        rr.font.size = Pt(10)

    out = OUT_DIR / spec["filename"]
    try:
        doc.save(out)
    except PermissionError:
        out = OUT_DIR / "claims-manager-work-hours-invoice-2026-08-20-more0.docx"
        doc.save(out)
        print(f"Target locked; wrote alternate {out}")
    else:
        print(f"Wrote {out}")
    print(f"  {spec['number']}: {len(entries)} entries, {total_h} h, payable {fmt_aud(total_inc)}")
    print(f"  period ledger: {entries[0]['date']} .. {entries[-1]['date']}")
    return out


def main() -> None:
    entries = parse_ledger(WORK_HOURS)
    build_invoice(SPEC, entries)


if __name__ == "__main__":
    main()
