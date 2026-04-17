"""
Generate a Microsoft Word invoice from docs/tracking/work_hours.md.
Usage: python scripts/generate-work-hours-invoice-docx.py [--date YYYY-MM-DD]
"""
from __future__ import annotations

import argparse
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

TABLE_CELL_PADDING_PT = 5.0


ROOT = Path(__file__).resolve().parents[1]
LEDGER = ROOT / "docs" / "tracking" / "work_hours.md"
INVOICES_DIR = ROOT / "docs" / "tracking" / "invoices"

VENDOR_NAME = "Kiama Hire Pty Ltd"
VENDOR_ABN = "61 671 383 136"
CUSTOMER_NAME = "Branlamie Pty Ltd"
CUSTOMER_ABN = "91 117 499 20"
CUSTOMER_ATTN = "Brett Elsey"

HOURLY_RATE_AUD = 125.0
GST_RATE = 0.10  # 10% GST on taxable supply (Australia)
BANK_BSB = "062 562"
BANK_ACCOUNT_NO = "1033 7743"

# Optional ledger line after metrics: "Lay summary: ..." (see work_hours_guide.md).
LAY_SUMMARY_LEDGER = re.compile(r"^Lay summary:\s*(.+)$", re.IGNORECASE)

# Optional trailing (NOT-BILLABLE) / (NOT_BILLABLE) on header (same meaning as | NOT_BILLABLE on metrics).
ENTRY_START = re.compile(
    r"^-\s+`(?P<d>\d{4}-\d{2}-\d{2})`\s+`(?P<sha>[0-9a-fA-F]+)`\s+"
    r"\*\*(?P<hours>[\d.]+)\s+h\*\*"
    r"(?:\s*\(\s*(?P<hdr_nb>NOT[-_\s]*BILLABLE)\s*\))?\s*$",
    re.IGNORECASE,
)

# Inline markdown: backticks first so content like `**x**` stays literal.
INLINE_MD = re.compile(r"(`[^`]+`|\*\*.+?\*\*)")

SUMMARY_MAX_CHARS = 280  # single short sentence for non-technical readers


def _strip_inline_markdown(text: str) -> str:
    t = re.sub(r"\*\*(.+?)\*\*", r"\1", text, flags=re.DOTALL)
    t = re.sub(r"`([^`]+)`", r"\1", t)
    return re.sub(r"\s+", " ", t).strip()


def _simplify_conventional_commit_lead(s: str) -> str:
    """Drop feat(scope): style prefixes for a plainer one-line lead."""
    t = re.sub(
        r"^(feat|fix|chore|refactor|docs|style|test|perf|ci|build)(\([^)]+\))?:\s*",
        "",
        s,
        flags=re.IGNORECASE,
    ).strip()
    return t or s


def _sentence_case(text: str) -> str:
    t = text.strip()
    if len(t) > 1 and t[0].isalpha() and t[0].islower():
        return t[0].upper() + t[1:]
    return t


def _cap_summary(text: str, max_chars: int = SUMMARY_MAX_CHARS) -> str:
    text = text.strip()
    if len(text) <= max_chars:
        return text
    cut = text[: max_chars - 1]
    sp = cut.rfind(" ")
    if sp > max_chars // 2:
        cut = cut[:sp]
    return cut + "…"


def _dejargon_for_lay_audience(text: str) -> str:
    """Swap common technical terms for plainer wording (best-effort, no LLM)."""
    t = text
    subs: list[tuple[str, str]] = [
        (r"\binfrastructure[- ]as[- ]code\b", "repeatable hosting setup"),
        (r"\bKubernetes\b|\bK8s\b", "cloud hosting"),
        (r"\bTerraform\b", "hosting templates"),
        (r"\bDocker\b", "packaged runtimes"),
        (r"\bNext\.js\b|\bNextJS\b", "the main website"),
        (r"\bmonorepo\b", "shared codebase"),
        (r"\bJWT\b", "secure sign-in"),
        (r"\bDrizzle\b", "database tools"),
        (r"\bOAuth\b|\bSSO\b", "industry-standard login"),
        (r"\bwebhooks?\b", "automatic notifications"),
        (r"\bAPI endpoints\b", "system connections"),
        (r"\bAPI layer\b", "system layer"),
        (r"\bAPIs?\b", "system services"),
        (r"\bCRUD\b", "add-and-edit screens"),
        (r"\bmigrations?\b", "database upgrades"),
        (r"\bschema(s)?\b", "data layout"),
        (r"\bfrontend\b", "user-facing website"),
        (r"\bbackend\b", "central system"),
        (r"\bNestJS\b|\bNest\b", "server software"),
        (r"\bmulti-tenant\b", "separate for each customer"),
        (r"\btenant(s)?\b", "customer organization"),
        (r"\bintegration schema\b", "how systems share data"),
        (r"\bintegration(s)?\b", "system links"),
        (r"\b\.gitignore\b", "ignore rules for build files"),
        (r"\bcompiled artifacts\b", "build output files"),
        (r"\bversion control\b", "change history"),
        (r"\bsquashed\b", "combined"),
        (r"\bbaseline\b", "starting snapshot"),
        (r"\bJWT organization_id\b", "organization on sign-in"),
        (r"\borganization_id\b", "organization marker"),
        (r"\brepository\b", "project files"),
        (r"\bpost-commit hook\b", "after-save script"),
        (r"\bcommit-level\b", "per-change"),
        (r"\bledger\b", "hours log"),
        (r"\bUI component library\b", "reusable screen parts"),
        (r"\bUI\b", "screen"),
        (r"\btyped client\b", "typed data link"),
        (r"\bcontainer images\b", "packaged app copies"),
        (r"\boperator scripts\b", "runbooks for staff"),
        (r"\bobservability\b", "monitoring"),
    ]
    for pat, rep in subs:
        t = re.sub(pat, rep, t, flags=re.IGNORECASE)
    return re.sub(r"\s+", " ", t).strip()


def _first_sentence(text: str) -> str:
    text = text.strip()
    if not text:
        return text
    parts = re.split(r"(?<=[.!?])\s+", text)
    return parts[0].strip() if parts else text


def description_to_lay_summary(description: str) -> str:
    """
    One plain-English sentence for a non-technical reader, from the invoice
    description (markdown stripped, jargon softened, first sentence only).
    """
    raw = description.strip()
    if not raw:
        return "No work summary was recorded for this line item."

    plain = _strip_inline_markdown(raw)
    plain = re.sub(r"\s+", " ", plain).strip()

    first_line = raw.split("\n", 1)[0].strip()
    m_lead = re.match(r"^\s*\*\*(.+?)\*\*(.*)$", first_line, flags=re.DOTALL)
    if m_lead:
        lead = re.sub(r"\s+", " ", m_lead.group(1).strip())
        lead = _simplify_conventional_commit_lead(lead)
        tail = m_lead.group(2).strip()
        if tail:
            lead = (lead + " " + _strip_inline_markdown(tail)).strip()
            lead = re.sub(r"\s+", " ", lead)
        candidate = lead if len(lead) >= 35 else _first_sentence(plain) or plain
    else:
        candidate = _first_sentence(plain) or plain

    candidate = _dejargon_for_lay_audience(candidate)
    one = _first_sentence(candidate) or candidate
    one = _sentence_case(one.strip())
    if one and not one.endswith((".", "!", "?")):
        one += "."
    return _cap_summary(one)


def _segment_is_non_billable(part: str) -> bool:
    """True for NOT_BILLABLE, NOT-BILLABLE, NOT BILLABLE, etc."""
    normalized = re.sub(r"[-\s]+", "_", part.strip().casefold())
    return normalized == "not_billable"


def split_not_billable_metrics(inner: str) -> tuple[str, bool]:
    """
    If the metrics line includes a pipe segment NOT_BILLABLE / NOT-BILLABLE (case-insensitive),
    return metrics with that segment removed and non_billable True.
    """
    raw = inner.strip()
    if not raw:
        return "", False
    parts = [p.strip() for p in raw.split("|")]
    is_non_billable = any(_segment_is_non_billable(p) for p in parts)
    kept = [p for p in parts if not _segment_is_non_billable(p)]
    return " | ".join(kept).strip(), is_non_billable


def parse_ledger(text: str) -> list[dict[str, str | bool]]:
    lines = text.splitlines()
    entries: list[dict[str, str | bool]] = []
    i = 0
    while i < len(lines):
        m = ENTRY_START.match(lines[i])
        if not m:
            i += 1
            continue
        date_s = m.group("d")
        sha = m.group("sha")
        hours = m.group("hours")
        header_non_billable = m.group("hdr_nb") is not None
        i += 1
        metrics = ""
        metrics_non_billable = False
        if i < len(lines):
            stripped = lines[i].strip()
            if stripped.startswith("`") and stripped.endswith("`"):
                inner = stripped[1:-1]
                metrics, metrics_non_billable = split_not_billable_metrics(inner)
                i += 1
        lay_summary = ""
        if i < len(lines):
            st_ls = lines[i].strip()
            m_ls = LAY_SUMMARY_LEDGER.match(st_ls)
            if m_ls:
                lay_summary = m_ls.group(1).strip()
                i += 1
        billable = not (header_non_billable or metrics_non_billable)
        desc_lines: list[str] = []
        while i < len(lines):
            line = lines[i]
            if ENTRY_START.match(line):
                break
            if line.strip() == "":
                if i + 1 < len(lines) and ENTRY_START.match(lines[i + 1]):
                    i += 1
                    break
                desc_lines.append("")
                i += 1
                continue
            desc_lines.append(line.strip())
            i += 1
        description = "\n".join(desc_lines).strip()
        entries.append(
            {
                "date": date_s,
                "commit": sha,
                "hours": hours,
                "metrics": metrics,
                "lay_summary": lay_summary,
                "description": description,
                "billable": billable,
            }
        )
    return entries


def _set_arial(run) -> None:
    run.font.name = "Arial"
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for tag in (qn("w:ascii"), qn("w:hAnsi"), qn("w:cs"), qn("w:eastAsia")):
        rfonts.set(tag, "Arial")


def set_document_arial_styles(doc: Document) -> None:
    """Default paragraph and common heading styles use Arial."""
    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
        try:
            st = doc.styles[style_name]
        except KeyError:
            continue
        st.font.name = "Arial"
        if style_name == "Normal" and st.font.size is None:
            st.font.size = Pt(11)


def set_cell_text(cell, text: str, *, bold: bool = False, size_pt: float | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    _set_arial(run)


def _apply_run_size(run, size_pt: float | None) -> None:
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    _set_arial(run)


def add_markdown_runs(paragraph, text: str, *, size_pt: float | None = None) -> None:
    """Turn `` `code` `` into italic Arial and **bold** into Word bold; plain text Arial."""
    pos = 0
    for m in INLINE_MD.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos : m.start()])
            _apply_run_size(run, size_pt)
        chunk = m.group(0)
        if chunk.startswith("`"):
            run = paragraph.add_run(chunk[1:-1])
            run.italic = True
            _apply_run_size(run, size_pt)
        else:
            run = paragraph.add_run(chunk[2:-2])
            run.bold = True
            _apply_run_size(run, size_pt)
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        _apply_run_size(run, size_pt)


def set_cell_markdown(cell, text: str, *, size_pt: float = 9) -> None:
    """Multi-line description: one Word paragraph per ledger line; inline markdown formatted."""
    cell.text = ""
    lines = text.split("\n")
    for idx, line in enumerate(lines):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        if not line.strip():
            continue
        add_markdown_runs(p, line, size_pt=size_pt)


def set_cell_entry_block(cell, entry: dict[str, str | bool], *, size_pt: float = 9) -> None:
    """Single cell: dash line then Date / Hours / Metrics / Commit (labels bold)."""
    cell.text = ""
    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_after = Pt(0)
    dash = p0.add_run("-")
    _apply_run_size(dash, size_pt)

    if entry.get("billable") is False:
        pnb = cell.add_paragraph()
        pnb.paragraph_format.space_after = Pt(2)
        rnb = pnb.add_run("NON-BILLABLE")
        rnb.bold = True
        _apply_run_size(rnb, size_pt)

    def add_label_paragraph(
        label: str,
        *,
        md_value: str | None = None,
        plain_value: str | None = None,
        space_before_pt: float | None = None,
    ) -> None:
        p = cell.add_paragraph()
        if space_before_pt is not None:
            p.paragraph_format.space_before = Pt(space_before_pt)
        p.paragraph_format.space_after = Pt(0)
        rl = p.add_run(label)
        rl.bold = True
        _apply_run_size(rl, size_pt)
        if plain_value is not None:
            rr = p.add_run(plain_value)
            _apply_run_size(rr, size_pt)
        elif md_value is not None:
            add_markdown_runs(p, md_value, size_pt=size_pt)

    add_label_paragraph("Date: ", plain_value=entry["date"])
    add_label_paragraph("Hours: ", plain_value=f'{entry["hours"]} h')
    add_label_paragraph("Metrics: ", md_value=entry["metrics"])
    p_commit = cell.add_paragraph()
    p_commit.paragraph_format.space_after = Pt(0)
    rl = p_commit.add_run("Commit: ")
    rl.bold = True
    _apply_run_size(rl, size_pt)
    rc = p_commit.add_run(entry["commit"])
    _apply_run_size(rc, size_pt)

    ledger_summary = str(entry.get("lay_summary") or "").strip()
    summary = (
        ledger_summary
        if ledger_summary
        else description_to_lay_summary(str(entry.get("description", "")))
    )
    add_label_paragraph("Summary: ", plain_value=summary, space_before_pt=6)


def add_invoice_header(doc: Document, *, invoice_date: str) -> None:
    """Vendor / customer blocks and invoice meta (first page)."""
    invoice_no = f"WH-{invoice_date}"

    tbl = doc.add_table(rows=2, cols=2)
    tbl.autofit = False
    # Portrait Letter usable width ~6.5"; split vendor / invoice meta.
    tbl.columns[0].width = Inches(3.35)
    tbl.columns[1].width = Inches(3.15)

    vendor_cell = tbl.rows[0].cells[0]
    meta_cell = tbl.rows[0].cells[1]

    p = vendor_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    rv = p.add_run(VENDOR_NAME)
    rv.bold = True
    rv.font.size = Pt(14)
    _set_arial(rv)
    p2 = vendor_cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    rv2 = p2.add_run(f"ABN: {VENDOR_ABN}")
    rv2.font.size = Pt(10)
    _set_arial(rv2)

    pm = meta_cell.paragraphs[0]
    pm.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pm.paragraph_format.space_after = Pt(2)
    rm = pm.add_run("INVOICE")
    rm.bold = True
    rm.font.size = Pt(22)
    _set_arial(rm)
    pm2 = meta_cell.add_paragraph()
    pm2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rm2 = pm2.add_run(f"Invoice date: {invoice_date}")
    rm2.font.size = Pt(10)
    _set_arial(rm2)
    pm3 = meta_cell.add_paragraph()
    pm3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rm3 = pm3.add_run(f"Invoice no.: {invoice_no}")
    rm3.font.size = Pt(10)
    _set_arial(rm3)

    left = tbl.rows[1].cells[0]
    right = tbl.rows[1].cells[1]
    bill_cell = left.merge(right)
    p = bill_cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    rl = p.add_run("Bill to")
    rl.bold = True
    rl.font.size = Pt(10)
    _set_arial(rl)
    for line in (
        CUSTOMER_NAME,
        f"ABN {CUSTOMER_ABN}",
        f"ATTN: {CUSTOMER_ATTN}",
    ):
        pl = bill_cell.add_paragraph()
        pl.paragraph_format.space_after = Pt(0)
        rx = pl.add_run(line)
        rx.font.size = Pt(10)
        _set_arial(rx)

    apply_table_cell_padding_pt(tbl, TABLE_CELL_PADDING_PT)


def set_cell_padding_pt(cell, padding_pt: float) -> None:
    """Inside margin for a table cell (dxa twips; improves readability)."""
    twips = max(1, int(round(float(padding_pt) * 20)))
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    old = tc_pr.find(qn("w:tcMar"))
    if old is not None:
        tc_pr.remove(old)
    tc_mar = OxmlElement("w:tcMar")
    for side in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(twips))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def apply_table_cell_padding_pt(table, padding_pt: float = TABLE_CELL_PADDING_PT) -> None:
    for row in table.rows:
        for cell in row.cells:
            set_cell_padding_pt(cell, padding_pt)


def set_document_portrait(doc: Document) -> None:
    """US Letter portrait (default reading layout)."""
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)


def _format_aud(amount: float) -> str:
    return f"${amount:,.2f}"


def add_invoice_amounts_and_payment(doc: Document, *, billable_hours: float) -> None:
    """Billable hours × rate + GST, then pay-to and bank details."""
    sub_ex = round(float(billable_hours) * HOURLY_RATE_AUD, 2)
    gst_amt = round(sub_ex * GST_RATE, 2)
    total_inc = round(sub_ex + gst_amt, 2)

    doc.add_paragraph()

    h = doc.add_heading("Amount & payment", level=2)
    for run in h.runs:
        _set_arial(run)

    rows: list[tuple[str, bool]] = [
        (f"Billable hours: {billable_hours:g} h", False),
        (f"Rate: {_format_aud(HOURLY_RATE_AUD)}/hr (ex GST)", False),
        (f"Subtotal (ex GST): {_format_aud(sub_ex)}", False),
        (f"GST ({GST_RATE:.0%}): {_format_aud(gst_amt)}", False),
        (f"Total payable (inc GST): {_format_aud(total_inc)}", True),
    ]
    for text, is_bold in rows:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = is_bold
        r.font.size = Pt(11 if is_bold else 10)
        _set_arial(r)

    doc.add_paragraph()
    p_pay = doc.add_paragraph()
    r_l = p_pay.add_run("Pay to: ")
    r_l.bold = True
    r_l.font.size = Pt(10)
    _set_arial(r_l)
    r_n = p_pay.add_run(VENDOR_NAME)
    r_n.font.size = Pt(10)
    _set_arial(r_n)

    for line in (f"BSB: {BANK_BSB}", f"Account: {BANK_ACCOUNT_NO}"):
        pl = doc.add_paragraph()
        rx = pl.add_run(line)
        rx.font.size = Pt(10)
        _set_arial(rx)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--date",
        default=None,
        help="Invoice date YYYY-MM-DD (default: today UTC-style from system)",
    )
    parser.add_argument(
        "--output",
        default=None,
        help="Output .docx path (default: docs/tracking/invoices/claims-manager-work-hours-invoice-{date}.docx)",
    )
    args = parser.parse_args()
    invoice_date = args.date or date.today().isoformat()

    text = LEDGER.read_text(encoding="utf-8")
    entries = parse_ledger(text)
    if not entries:
        raise SystemExit(f"No entries parsed from {LEDGER}")

    billable_total = sum(float(e["hours"]) for e in entries if e["billable"] is not False)
    non_billable_total = sum(float(e["hours"]) for e in entries if e["billable"] is False)
    non_billable_count = sum(1 for e in entries if e["billable"] is False)

    doc = Document()
    set_document_arial_styles(doc)
    set_document_portrait(doc)
    add_invoice_header(doc, invoice_date=invoice_date)

    doc.add_paragraph()

    sec = doc.add_heading("Work hours line items", level=1)
    for run in sec.runs:
        _set_arial(run)

    sub = doc.add_paragraph()
    r = sub.add_run(
        "Project: claims-manager\n"
        "Basis: docs/tracking/work_hours.md (commit-level ledger)\n"
        f"Invoice date: {invoice_date}"
    )
    r.font.size = Pt(10)
    _set_arial(r)

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    # Portrait: entry + description within ~6.5" usable width.
    widths = (Inches(2.25), Inches(4.25))
    hdr = table.rows[0].cells
    headers = ("Entry", "Description (invoice text)")
    for cell, label, w in zip(hdr, headers, widths):
        set_cell_text(cell, label, bold=True, size_pt=10)
        cell.width = w

    for e in entries:
        row = table.add_row().cells
        set_cell_entry_block(row[0], e, size_pt=9)
        set_cell_markdown(row[1], e["description"], size_pt=9)

    total_row = table.add_row().cells
    set_cell_text(total_row[0], f"Total (billable): {billable_total:g} h", bold=True, size_pt=10)
    if non_billable_total > 0:
        p2 = total_row[0].add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(f"Non-billable (excluded from total): {non_billable_total:g} h")
        r2.bold = True
        r2.font.size = Pt(9)
        _set_arial(r2)
    right_total = f"{len(entries)} line items"
    if non_billable_count:
        right_total += f"\n{non_billable_count} non-billable (hours excluded from billable total)"
    set_cell_text(total_row[1], right_total, bold=True, size_pt=10)

    apply_table_cell_padding_pt(table, TABLE_CELL_PADDING_PT)

    add_invoice_amounts_and_payment(doc, billable_hours=billable_total)

    INVOICES_DIR.mkdir(parents=True, exist_ok=True)
    out = (
        Path(args.output).resolve()
        if args.output
        else INVOICES_DIR / f"claims-manager-work-hours-invoice-{invoice_date}.docx"
    )
    try:
        doc.save(out)
    except PermissionError:
        if not args.output:
            alt = out.with_name(out.stem + "-pending.docx")
            doc.save(alt)
            print(
                f"Could not overwrite {out} (file may be open). "
                f"Wrote {alt} instead. Close the document and re-run, or use --output."
            )
            return
        raise
    print(
        f"Wrote {out} ({len(entries)} line items, "
        f"billable {billable_total:g} h"
        + (f", non-billable {non_billable_total:g} h excluded" if non_billable_total else "")
        + ")"
    )


if __name__ == "__main__":
    main()
