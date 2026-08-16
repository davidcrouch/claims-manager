"""Build claims-manager work-hours invoices (cash-flow plan: 2026-07-05 and 2026-08-16).

Consistent with WH-2026-04-14 / WH-2026-04-30:
- Kiama Hire Pty Ltd -> Branlamie Pty Ltd
- $125.00/hr ex GST + 10% GST
- Commit-level line items from docs/tracking/work_hours.md
"""
from __future__ import annotations

import re
from decimal import Decimal
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

A = "Arial"
WORK_HOURS = Path(__file__).resolve().parents[1] / "work_hours.md"
OUT_DIR = Path(__file__).resolve().parent
RATE_H = Decimal("125.00")

# Cash-flow alternative commit sets (whole commits, chronological).
INV_2026_07_05_SHAS = {
    "ddc2b5d",
    "820d7a2",
    "8269b5f",
    "ddc250a",
    "a78f00b",
    "806409c",
    "5c38559",
    "96497e6",
    "e50a92f",
    "9393b56",
}

INV_2026_08_16_SHAS = {
    "1a6026f",
    "a1c3eb8",
    "f92bd46",
    "7e162e2",
    "a9801e4",
    "dd4589c",
    "0b10477",
    "c48cdab",
    "49eeca5",
    "5fd5eb6",
    "c83063c",
    "c08fc25",
    "b676b82",
    "187b34a",
    "4ecff44",
    "d875fa9",
    "2e9e0df",
    "038dff4",
    "0f2856a",
    "55d15fd",
    "296e7e9",
    "da322cf",
    "d12a41f",
    "08f1bbd",
    "f8692c3",
    "43c5468",
    "69aa3c3",
}

INVOICES = [
    {
        "date": "2026-07-05",
        "number": "WH-2026-07-05",
        "filename": "claims-manager-work-hours-invoice-2026-07-05.docx",
        "shas": INV_2026_07_05_SHAS,
        "period": (
            "Period: new billable ledger entries after invoice WH-2026-04-30 "
            "(entries dated 2026-05-13 through 2026-06-30)."
        ),
        "prior": "WH-2026-04-30",
    },
    {
        "date": "2026-08-16",
        "number": "WH-2026-08-16",
        "filename": "claims-manager-work-hours-invoice-2026-08-16.docx",
        "shas": INV_2026_08_16_SHAS,
        "period": (
            "Period: new billable ledger entries after invoice WH-2026-07-05 "
            "(entries dated 2026-07-30 through 2026-08-09)."
        ),
        "prior": "WH-2026-07-05",
    },
]


def parse_ledger(path: Path) -> list[dict]:
    content = path.read_text(encoding="utf-8")
    header_re = re.compile(
        r"^- `(?P<date>\d{4}-\d{2}-\d{2})` `(?P<commit>[a-f0-9]+)` \*\*(?P<hours>[\d.]+) h\*\*(?P<suffix>[^\n]*)",
        re.MULTILINE,
    )
    matches = list(header_re.finditer(content))
    out: list[dict] = []
    for i, m in enumerate(matches):
        if "NOT-BILLABLE" in m.group("suffix"):
            continue
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(content)
        block = content[start:end]
        lines = block.strip().split("\n")
        metrics = ""
        idx = 0
        if lines and lines[0].lstrip().startswith("`"):
            metrics = lines[0].strip().strip("`").strip()
            idx = 1
        lay = ""
        while idx < len(lines):
            s = lines[idx].strip()
            if s.startswith("Lay summary:"):
                lay = s.split("Lay summary:", 1)[1].strip()
                idx += 1
                break
            idx += 1
        desc: list[str] = []
        while idx < len(lines):
            raw = lines[idx].rstrip()
            if raw.strip().startswith("- `") and "` `" in raw and "**" in raw and " h**" in raw:
                break
            if raw.strip():
                desc.append(raw)
            idx += 1
        out.append(
            {
                "date": m.group("date"),
                "commit": m.group("commit"),
                "hours": m.group("hours") + " h",
                "hours_n": Decimal(m.group("hours")),
                "metrics": metrics,
                "lay": lay,
                "desc": desc,
            }
        )
    return out


def add_runs_from_md(paragraph, text: str, size_pt: float) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        body = part[2:-2] if bold else part
        r = paragraph.add_run(body)
        r.bold = bold
        r.font.name = A
        r.font.size = Pt(size_pt)


def fmt_aud(d: Decimal) -> str:
    return f"${d:,.2f}"


def build_invoice(spec: dict, ledger: list[dict]) -> Path:
    sha_set: set[str] = spec["shas"]
    # Preserve ledger order; match by short or full sha prefix used in ledger.
    entries = [e for e in ledger if e["commit"] in sha_set or any(e["commit"].startswith(s) for s in sha_set)]
    # Prefer exact short-sha match used in ledger (7 chars typically)
    entries = [e for e in ledger if e["commit"] in sha_set]
    missing = sha_set - {e["commit"] for e in entries}
    if missing:
        raise SystemExit(f"{spec['number']}: missing commits in ledger: {sorted(missing)}")
    if len(entries) != len(sha_set):
        raise SystemExit(
            f"{spec['number']}: expected {len(sha_set)} entries, got {len(entries)}"
        )

    total_h = sum((e["hours_n"] for e in entries), Decimal("0"))
    sub_ex = total_h * RATE_H
    gst = (sub_ex * Decimal("0.10")).quantize(Decimal("0.01"))
    total_inc = sub_ex + gst

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    ht = doc.add_table(rows=2, cols=2)
    c00, c01 = ht.rows[0].cells[0], ht.rows[0].cells[1]
    c10, c11 = ht.rows[1].cells[0], ht.rows[1].cells[1]
    c10.merge(c11)

    p = c00.paragraphs[0]
    r = p.add_run("Kiama Hire Pty Ltd")
    r.bold = True
    r.font.name = A
    r.font.size = Pt(14)
    p = c00.add_paragraph()
    r = p.add_run("ABN: 61 671 383 136")
    r.font.name = A
    r.font.size = Pt(10)

    p = c01.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("INVOICE")
    r.bold = True
    r.font.name = A
    r.font.size = Pt(22)
    p = c01.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"Invoice date: {spec['date']}")
    r.font.name = A
    r.font.size = Pt(10)
    p = c01.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"Invoice no.: {spec['number']}")
    r.font.name = A
    r.font.size = Pt(10)

    p = c10.paragraphs[0]
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run("Bill to")
    r.bold = True
    r.font.name = A
    r.font.size = Pt(10)
    for line in ("Branlamie Pty Ltd", "ABN 91 117 499 20", "ATTN: Brett Elsey"):
        p = c10.add_paragraph()
        r = p.add_run(line)
        r.font.name = A
        r.font.size = Pt(10)

    doc.add_paragraph()
    h = doc.add_heading("Work hours line items", level=1)
    for run in h.runs:
        run.font.name = A

    intro = doc.add_paragraph()
    bits = [
        "Project: claims-manager",
        "Basis: docs/tracking/work_hours.md (commit-level ledger)",
        spec["period"],
        f"Invoice date: {spec['date']}",
    ]
    for i, line in enumerate(bits):
        if i:
            intro.add_run().add_break()
        rr = intro.add_run(line)
        rr.font.name = A
        rr.font.size = Pt(10)

    doc.add_paragraph()

    tbl = doc.add_table(rows=1 + len(entries) + 1, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for j, title in enumerate(("Entry", "Description (invoice text)")):
        rr = hdr[j].paragraphs[0].add_run(title)
        rr.bold = True
        rr.font.name = A
        rr.font.size = Pt(10)

    for i, e in enumerate(entries):
        row = tbl.rows[i + 1].cells
        left, right = row[0], row[1]
        left.text = ""
        lp = left.paragraphs[0]
        r = lp.add_run("-")
        r.font.name = A
        r.font.size = Pt(9)

        def left_line(label: str, value: str) -> None:
            p = left.add_paragraph()
            lr = p.add_run(f"{label}: ")
            lr.bold = True
            lr.font.name = A
            lr.font.size = Pt(9)
            vr = p.add_run(value)
            vr.font.name = A
            vr.font.size = Pt(9)

        left_line("Date", e["date"])
        left_line("Hours", e["hours"])
        left_line("Metrics", e["metrics"])
        left_line("Commit", e["commit"])
        sp = left.add_paragraph()
        sp.paragraph_format.space_before = Pt(6)
        sr = sp.add_run("Summary: ")
        sr.bold = True
        sr.font.name = A
        sr.font.size = Pt(9)
        sr2 = sp.add_run(e["lay"])
        sr2.font.name = A
        sr2.font.size = Pt(9)

        right.text = ""
        for j, line in enumerate(e["desc"]):
            rp = right.paragraphs[0] if j == 0 else right.add_paragraph()
            rp.paragraph_format.space_after = Pt(0)
            add_runs_from_md(rp, line, 9)

    tot = tbl.rows[-1].cells
    tot[0].text = ""
    tp = tot[0].paragraphs[0]
    tr = tp.add_run(f"Total (billable): {total_h} h")
    tr.bold = True
    tr.font.name = A
    tr.font.size = Pt(10)
    tot[1].text = ""
    tp2 = tot[1].paragraphs[0]
    tr2 = tp2.add_run(f"{len(entries)} line items")
    tr2.bold = True
    tr2.font.name = A
    tr2.font.size = Pt(10)

    doc.add_paragraph()
    h3 = doc.add_heading("Amount & payment", level=2)
    for run in h3.runs:
        run.font.name = A

    # Match WH-2026-04-14 style when there are no reimbursable Cursor charges.
    amount_lines = [
        (f"Billable hours: {total_h} h", False, 10),
        ("Rate: $125.00/hr (ex GST)", False, 10),
        (f"Subtotal (ex GST): {fmt_aud(sub_ex)}", False, 10),
        (f"GST (10%): {fmt_aud(gst)}", False, 10),
        (f"Total payable (inc GST): {fmt_aud(total_inc)}", True, 11),
    ]
    for text, bold, sz in amount_lines:
        p = doc.add_paragraph()
        rr = p.add_run(text)
        rr.bold = bold
        rr.font.name = A
        rr.font.size = Pt(sz)

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Pay to: ")
    r.bold = True
    r.font.name = A
    r.font.size = Pt(10)
    r2 = p.add_run("Kiama Hire Pty Ltd")
    r2.font.name = A
    r2.font.size = Pt(10)
    for line in ("BSB: 062 562", "Account: 1033 7743"):
        p = doc.add_paragraph()
        rr = p.add_run(line)
        rr.font.name = A
        rr.font.size = Pt(10)

    out = OUT_DIR / spec["filename"]
    doc.save(out)
    print(f"Wrote {out}")
    print(f"  {spec['number']}: {len(entries)} entries, {total_h} h, payable {fmt_aud(total_inc)}")
    return out


def main() -> None:
    ledger = parse_ledger(WORK_HOURS)
    for spec in INVOICES:
        build_invoice(spec, ledger)


if __name__ == "__main__":
    main()
