"""Rebuild claims-manager-work-hours-invoice-2026-04-30.docx from ledger + Cursor receipts."""
from __future__ import annotations

import re
from decimal import Decimal
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

A = "Arial"
WORK_HOURS = Path(__file__).resolve().parents[1] / "work_hours.md"
OUT = Path(__file__).resolve().parent / "claims-manager-work-hours-invoice-2026-04-30.docx"
AFTER_DATE = "2026-04-14"
RATE = Decimal("1.5055")

CURSOR = [
    ("2038-7192", "April 16, 2026", Decimal("22.73"), (Decimal("22.73") * RATE).quantize(Decimal("0.01"))),
    ("2443-8287", "March 31, 2026", Decimal("220.00"), Decimal("331.22")),
    ("2716-2237", "April 27, 2026", Decimal("60.46"), (Decimal("60.46") * RATE).quantize(Decimal("0.01"))),
    ("2806-2142", "April 30, 2026", Decimal("80.13"), (Decimal("80.13") * RATE).quantize(Decimal("0.01"))),
    ("2928-5652", "April 26, 2026", Decimal("44.51"), (Decimal("44.51") * RATE).quantize(Decimal("0.01"))),
]


def parse_ledger(path: Path, after: str) -> list[dict]:
    content = path.read_text(encoding="utf-8")
    header_re = re.compile(
        r"^- `(?P<date>\d{4}-\d{2}-\d{2})` `(?P<commit>[a-f0-9]+)` \*\*(?P<hours>[\d.]+) h\*\*(?P<suffix>[^\n]*)",
        re.MULTILINE,
    )
    matches = list(header_re.finditer(content))
    out: list[dict] = []
    for i, m in enumerate(matches):
        if "NOT-BILLABLE" in m.group("suffix"):
            continue
        date = m.group("date")
        if date <= after:
            continue
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(content)
        block = content[start:end]
        lines = block.strip().split("\n")
        metrics = ""
        desc: list[str] = []
        idx = 0
        if lines and lines[0].lstrip().startswith("`"):
            metrics = lines[0].strip().strip("`").strip()
            idx = 1
        while idx < len(lines):
            s = lines[idx].strip()
            if s.startswith("Lay summary:"):
                lay = s.split("Lay summary:", 1)[1].strip()
                idx += 1
                break
            idx += 1
        else:
            lay = ""
        while idx < len(lines):
            raw = lines[idx].rstrip()
            if raw.strip().startswith("- `") and "` `" in raw and "**" in raw and " h**" in raw:
                break
            if raw.strip():
                desc.append(raw)
            idx += 1
        out.append(
            {
                "date": date,
                "commit": m.group("commit"),
                "hours": m.group("hours") + " h",
                "metrics": metrics,
                "lay": lay,
                "desc": desc,
            }
        )
    return out


def add_runs_from_md(paragraph, text: str, size_pt: float) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        body = part[2:-2] if bold else part
        r = paragraph.add_run(body)
        r.bold = bold
        r.font.name = A
        r.font.size = Pt(size_pt)


def fmt_aud(d: Decimal) -> str:
    return f"${d:,.2f}"


def fmt_usd(d: Decimal) -> str:
    return f"${d:,.2f}"


def main() -> None:
    entries = parse_ledger(WORK_HOURS, AFTER_DATE)
    total_h = sum(Decimal(e["hours"].replace(" h", "")) for e in entries)
    exp_aud = sum(r[3] for r in CURSOR)
    rate_h = Decimal("125.00")
    sub_ex = total_h * rate_h
    gst = (sub_ex * Decimal("0.10")).quantize(Decimal("0.01"))
    prof_inc = sub_ex + gst
    grand = prof_inc + exp_aud

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    ht = doc.add_table(rows=2, cols=2)
    c00, c01 = ht.rows[0].cells[0], ht.rows[0].cells[1]
    c10, c11 = ht.rows[1].cells[0], ht.rows[1].cells[1]
    c10.merge(c11)

    p = c00.paragraphs[0]
    r = p.add_run("Kiama Hire Pty Ltd")
    r.bold = True
    r.font.name = A
    r.font.size = Pt(14)
    p = c00.add_paragraph()
    r = p.add_run("ABN: 61 671 383 136")
    r.font.name = A
    r.font.size = Pt(10)

    p = c01.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("INVOICE")
    r.bold = True
    r.font.name = A
    r.font.size = Pt(22)
    p = c01.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Invoice date: 2026-04-30")
    r.font.name = A
    r.font.size = Pt(10)
    p = c01.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Invoice no.: WH-2026-04-30")
    r.font.name = A
    r.font.size = Pt(10)

    p = c10.paragraphs[0]
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run("Bill to")
    r.bold = True
    r.font.name = A
    r.font.size = Pt(10)
    for line in ("Branlamie Pty Ltd", "ABN 91 117 499 20", "ATTN: Brett Elsey"):
        p = c10.add_paragraph()
        r = p.add_run(line)
        r.font.name = A
        r.font.size = Pt(10)

    doc.add_paragraph()
    h = doc.add_heading("Work hours line items", level=1)
    for run in h.runs:
        run.font.name = A

    intro = doc.add_paragraph()
    bits = [
        "Project: claims-manager",
        "Basis: docs/tracking/work_hours.md (commit-level ledger)",
        "Period: new billable ledger entries after invoice WH-2026-04-14 (entries dated 2026-04-17 through 2026-04-25; April remainder through month-end).",
        "Invoice date: 2026-04-30",
    ]
    for i, line in enumerate(bits):
        if i:
            intro.add_run().add_break()
        rr = intro.add_run(line)
        rr.font.name = A
        rr.font.size = Pt(10)

    doc.add_paragraph()

    tbl = doc.add_table(rows=1 + len(entries) + 1, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for j, title in enumerate(("Entry", "Description (invoice text)")):
        rr = hdr[j].paragraphs[0].add_run(title)
        rr.bold = True
        rr.font.name = A
        rr.font.size = Pt(10)

    for i, e in enumerate(entries):
        row = tbl.rows[i + 1].cells
        left, right = row[0], row[1]
        left.text = ""
        lp = left.paragraphs[0]
        r = lp.add_run("-")
        r.font.name = A
        r.font.size = Pt(9)

        def left_line(label: str, value: str) -> None:
            p = left.add_paragraph()
            lr = p.add_run(f"{label}: ")
            lr.bold = True
            lr.font.name = A
            lr.font.size = Pt(9)
            vr = p.add_run(value)
            vr.font.name = A
            vr.font.size = Pt(9)

        left_line("Date", e["date"])
        left_line("Hours", e["hours"])
        left_line("Metrics", e["metrics"])
        left_line("Commit", e["commit"])
        sp = left.add_paragraph()
        sp.paragraph_format.space_before = Pt(6)
        sr = sp.add_run("Summary: ")
        sr.bold = True
        sr.font.name = A
        sr.font.size = Pt(9)
        sr2 = sp.add_run(e["lay"])
        sr2.font.name = A
        sr2.font.size = Pt(9)

        right.text = ""
        for j, line in enumerate(e["desc"]):
            rp = right.paragraphs[0] if j == 0 else right.add_paragraph()
            rp.paragraph_format.space_after = Pt(0)
            add_runs_from_md(rp, line, 9)

    tot = tbl.rows[-1].cells
    tot[0].text = ""
    tp = tot[0].paragraphs[0]
    tr = tp.add_run(f"Total (billable): {total_h} h")
    tr.bold = True
    tr.font.name = A
    tr.font.size = Pt(10)
    tot[1].text = ""
    tp2 = tot[1].paragraphs[0]
    tr2 = tp2.add_run(f"{len(entries)} line items")
    tr2.bold = True
    tr2.font.name = A
    tr2.font.size = Pt(10)

    doc.add_paragraph()

    h2 = doc.add_heading("Third-party tooling (Cursor) - reimbursable", level=2)
    for run in h2.runs:
        run.font.name = A

    ex_p = doc.add_paragraph()
    er = ex_p.add_run(
        "Cursor subscription and usage receipts (USD) converted to Australian dollars at 1 USD = 1.5055 AUD. "
        "Receipt 2443-8287 AUD matches the processor charge (A$331.22); other rows use the same rate."
    )
    er.font.name = A
    er.font.size = Pt(10)

    et = doc.add_table(rows=1 + len(CURSOR) + 1, cols=5)
    et.style = "Table Grid"
    eh = ("Receipt", "Date paid", "Amount (USD)", "FX", "Amount (AUD)")
    for j, title in enumerate(eh):
        rr = et.rows[0].cells[j].paragraphs[0].add_run(title)
        rr.bold = True
        rr.font.name = A
        rr.font.size = Pt(10)
    for ri, (rec, dt, usd, aud) in enumerate(CURSOR, start=1):
        vals = (rec, dt, fmt_usd(usd), "1 USD = 1.5055 AUD", fmt_aud(aud))
        for j, val in enumerate(vals):
            rr = et.rows[ri].cells[j].paragraphs[0].add_run(val)
            rr.font.name = A
            rr.font.size = Pt(9)
    for j in range(4):
        et.rows[-1].cells[j].text = ""
    z = et.rows[-1].cells[0].paragraphs[0].add_run("Total reimbursable")
    z.bold = True
    z.font.name = A
    z.font.size = Pt(10)
    et.rows[-1].cells[4].text = ""
    zz = et.rows[-1].cells[4].paragraphs[0].add_run(fmt_aud(exp_aud))
    zz.bold = True
    zz.font.name = A
    zz.font.size = Pt(10)

    doc.add_paragraph()

    h3 = doc.add_heading("Amount & payment", level=2)
    for run in h3.runs:
        run.font.name = A

    lines = [
        (f"Billable hours: {total_h} h", False, 10),
        ("Rate: $125.00/hr (ex GST)", False, 10),
        (f"Subtotal (ex GST): {fmt_aud(sub_ex)}", False, 10),
        (f"GST (10%): {fmt_aud(gst)}", False, 10),
        (f"Total professional services (inc GST): {fmt_aud(prof_inc)}", True, 11),
        (f"Plus reimbursable Cursor charges (AUD): {fmt_aud(exp_aud)}", False, 10),
        (f"Grand total payable: {fmt_aud(grand)}", True, 11),
    ]
    for text, bold, sz in lines:
        p = doc.add_paragraph()
        rr = p.add_run(text)
        rr.bold = bold
        rr.font.name = A
        rr.font.size = Pt(sz)

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Pay to: ")
    r.bold = True
    r.font.name = A
    r.font.size = Pt(10)
    r2 = p.add_run("Kiama Hire Pty Ltd")
    r2.font.name = A
    r2.font.size = Pt(10)
    for line in ("BSB: 062 562", "Account: 1033 7743"):
        p = doc.add_paragraph()
        rr = p.add_run(line)
        rr.font.name = A
        rr.font.size = Pt(10)

    doc.save(OUT)
    print(f"Wrote {OUT}")
    print(f"Entries: {len(entries)}, hours: {total_h}, expenses AUD: {exp_aud}, grand: {grand}")


if __name__ == "__main__":
    main()
